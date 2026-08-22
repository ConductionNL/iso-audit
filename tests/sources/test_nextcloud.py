"""Tests voor `iso_audit.sources.nextcloud` — het Source-protocol op een WebDAV-bron.

De echte-server-check zit in `scripts/preflight.py --component nextcloud`; die draaide op
2026-08-22 tegen `canary-accept/nextcloud` 32.0.13 en vond twee fouten die deze gestubde tests
niet zagen. Wat hier staat is het gedragscontract: dekking die optelt, dezelfde lezers als
Drive, en niets dat stil verdwijnt.
"""

from __future__ import annotations

from typing import Any
from unittest.mock import patch

import pytest

from iso_audit.clients import nextcloud as dav
from iso_audit.sources import nextcloud as bron_mod
from iso_audit.sources.base import Document
from iso_audit.sources.tekst import LeegDocumentError


@pytest.fixture(autouse=True)
def _omgeving(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("NEXTCLOUD_BASE_URL", "https://cloud.example.org")
    monkeypatch.setenv("NEXTCLOUD_USER", "auditor")
    monkeypatch.setenv("NEXTCLOUD_APP_PASSWORD", "app-wachtwoord")
    monkeypatch.delenv("NEXTCLOUD_PATHS", raising=False)


DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _item(pad: str, mime: str, is_map: bool = False) -> dav.Item:
    return dav.Item(pad=pad, is_map=is_map, mime=mime, gewijzigd="", bytes_groot=1)


def test_zonder_configuratie_geen_stille_bron(monkeypatch: pytest.MonkeyPatch) -> None:
    """Een bron die zonder configuratie 'werkt' leest niets en meldt dat niet."""
    monkeypatch.delenv("NEXTCLOUD_APP_PASSWORD", raising=False)
    with pytest.raises(OSError, match="niet geconfigureerd"):
        bron_mod.NextcloudSource()


def test_list_documents_dekking_telt_op() -> None:
    """`gezien = gelezen + overgeslagen` is de rekensom die een auditor maakt.

    Op 2026-08-22 klopte die niet: wat de client al oversloeg (verborgen bestanden) werd wel
    in `overgeslagen` geteld maar niet in `gezien`.
    """
    items = [
        _item(
            "Audit/Beleid.docx",
            bron_mod.DOCX_MIME
            if hasattr(bron_mod, "DOCX_MIME")
            else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        _item("Audit/Foto.png", "image/png"),
        _item("Audit/Raadsel.bin", "application/x-onbekend"),
    ]

    def _lijst(_verb: Any, _pad: str, **kw: Any) -> list[dav.Item]:
        kw["overgeslagen"]["verborgen bestand of map"] = 2
        return items

    with patch.object(dav, "lijst_recursief", _lijst):
        bron = bron_mod.NextcloudSource(paden=["Audit"])
        docs = list(bron.list_documents())

    dekking = bron.dekking()
    assert [d.titel for d in docs] == ["Beleid.docx"]
    assert dekking["gezien"] == dekking["gelezen"] + sum(dekking["overgeslagen"].values())
    assert dekking["gezien"] == 5


def test_onbekend_type_en_niet_leesbaar_worden_gemeld() -> None:
    """Niets verdwijnt stil — de regel uit `landschap-dekking`, ook bij een nieuwe bron."""

    def _lijst(_verb: Any, _pad: str, **kw: Any) -> list[dav.Item]:
        return [_item("a/Foto.png", "image/png"), _item("a/X.bin", "application/x-onbekend")]

    with patch.object(dav, "lijst_recursief", _lijst):
        bron = bron_mod.NextcloudSource(paden=["a"])
        list(bron.list_documents())

    redenen = bron.dekking()["overgeslagen"]
    assert any("image/png" in r and "OCR" in r for r in redenen)
    assert any("onbekend type: application/x-onbekend" in r for r in redenen)


def test_dedup_op_pad() -> None:
    """Een bestand onder twee geconfigureerde paden telt één keer."""

    def _lijst(_verb: Any, pad: str, **kw: Any) -> list[dav.Item]:
        return [_item("a/b/Doc.txt", "text/plain")]

    with patch.object(dav, "lijst_recursief", _lijst):
        bron = bron_mod.NextcloudSource(paden=["a", "a/b"])
        docs = list(bron.list_documents())

    assert len(docs) == 1


def test_fetch_content_gebruikt_de_gedeelde_lezers() -> None:
    """Dezelfde docx-lezer als Drive, inclusief tabellen."""
    import io

    import docx as _docx

    d = _docx.Document()
    tabel = d.add_table(rows=1, cols=2)
    tabel.cell(0, 0).text = "Maatregel"
    tabel.cell(0, 1).text = "MFA"
    buffer = io.BytesIO()
    d.save(buffer)

    doc = Document(
        id="a/Beleid.docx",
        titel="Beleid.docx",
        bron="nextcloud",
        type="docx",
        laatst_gewijzigd="",
        inhoud_uri="a/Beleid.docx",
    )
    with patch.object(dav, "download", return_value=buffer.getvalue()):
        tekst = bron_mod.NextcloudSource().fetch_content(doc)

    assert "MFA" in tekst, "tabelcellen komen mee, net als bij Drive"


def test_leeg_bestand_is_een_storing_met_een_passende_reden() -> None:
    doc = Document(
        id="a/Leeg.txt",
        titel="Leeg.txt",
        bron="nextcloud",
        type="txt",
        laatst_gewijzigd="",
        inhoud_uri="a/Leeg.txt",
    )
    with (
        patch.object(dav, "download", return_value=b"   "),
        pytest.raises(LeegDocumentError, match="het bestand is leeg"),
    ):
        bron_mod.NextcloudSource().fetch_content(doc)


def test_document_uit_andere_bron_wordt_geweigerd() -> None:
    doc = Document(id="x", titel="x", bron="drive", type="txt", laatst_gewijzigd="", inhoud_uri="x")
    with pytest.raises(ValueError, match="verwacht 'nextcloud'"):
        bron_mod.NextcloudSource().fetch_content(doc)


def test_probe_geeft_status_per_pad() -> None:
    """Eén verkeerd pad mag een werkende configuratie niet als kapot laten ogen."""

    def _bereikbaar(_verb: Any, pad: str) -> tuple[bool, str]:
        return (pad == "goed", "ok" if pad == "goed" else "Dit pad bestaat niet op de server.")

    with patch.object(dav, "bereikbaar", _bereikbaar):
        uit = bron_mod.NextcloudSource(paden=["goed", "fout"]).probe()

    assert uit["status"] == "ok", "één werkend pad volstaat"
    statussen = {loc["naam"]: loc["status"] for loc in uit["locaties"]}
    assert statussen == {"goed": "ok", "fout": "fail"}


def test_probe_meldt_de_reden_als_niets_werkt() -> None:
    with patch.object(dav, "bereikbaar", lambda *a: (False, "Aanmelden mislukte")):
        uit = bron_mod.NextcloudSource(paden=["a"]).probe()
    assert uit["status"] == "fail"
    assert "Aanmelden mislukte" in str(uit["reden"])


def test_nextcloud_staat_in_de_registry() -> None:
    from iso_audit.ingest import beschikbare_bronnen

    assert "nextcloud" in beschikbare_bronnen()


def test_bron_url_wijst_naar_de_server(monkeypatch: pytest.MonkeyPatch) -> None:
    """Zonder link is 'open het document zelf' niet te doen."""
    from iso_audit.api.run_job import _bron_url

    monkeypatch.setenv("NEXTCLOUD_BASE_URL", "https://cloud.example.org")
    url = _bron_url("nextcloud", "Audit/Sub/Beleid.docx")
    assert url is not None and url.startswith("https://cloud.example.org")
    assert "Audit/Sub" in url
